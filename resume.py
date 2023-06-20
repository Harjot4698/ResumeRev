from flask import Flask, render_template, request, send_file
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
import requests
import os

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        url = request.form.get('url')
        file = request.files['file']
        filename = file.filename

        file.save(filename)

        text = scrape_website(url)
        write_to_word_file(text, filename)

        return send_file(filename, as_attachment=True)

    return render_template('index.html')

def scrape_website(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    return soup.get_text()

def write_to_word_file(text, path):
    doc = Document(path)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.color.rgb = RGBColor(255, 255, 255)  # white color
    font.size = Pt(1)  # minimal font size
    doc.save(path)

if __name__ == "__main__":
    app.run(port=5000, debug=True)
